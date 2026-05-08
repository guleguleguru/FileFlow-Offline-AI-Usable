from __future__ import annotations

from pathlib import Path

import fitz
import pytest
from docx import Document
from PIL import Image

from offline_converter import converters
from offline_converter.converters import (
    ConversionError,
    MissingDependencyError,
    image_to_pdf,
    pdf_to_images,
    pdf_to_word,
    word_to_pdf,
)


def make_image(path: Path, color: tuple[int, int, int], size: tuple[int, int] = (120, 80)) -> None:
    Image.new("RGB", size, color).save(path)


def test_image_to_pdf_combines_multiple_images(tmp_path: Path) -> None:
    first = tmp_path / "first.png"
    second = tmp_path / "second.jpg"
    make_image(first, (220, 40, 40))
    make_image(second, (40, 80, 220), size=(80, 120))

    output = tmp_path / "combined.pdf"
    result = image_to_pdf([first, second], output, quality=85)

    assert result.output_path == output
    assert result.page_count == 2
    with fitz.open(output) as pdf:
        assert pdf.page_count == 2


def test_pdf_to_images_exports_selected_pages(tmp_path: Path) -> None:
    images = []
    for index in range(3):
        image_path = tmp_path / f"page-{index + 1}.png"
        make_image(image_path, (index * 40, 100, 180))
        images.append(image_path)
    source_pdf = tmp_path / "source.pdf"
    Image.open(images[0]).save(source_pdf, save_all=True, append_images=[Image.open(p) for p in images[1:]])

    result = pdf_to_images(source_pdf, tmp_path / "pages", image_format="png", dpi=72, pages=[1, 3])

    assert result.page_count == 2
    assert [path.name for path in result.outputs] == ["source-page-001.png", "source-page-003.png"]
    assert all(path.exists() for path in result.outputs)


def test_pdf_to_word_extracts_editable_text_from_text_pdf(tmp_path: Path) -> None:
    source_pdf = tmp_path / "text.pdf"
    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((72, 72), "Hello PDF")
    chinese_font = Path("C:/Windows/Fonts/simhei.ttf")
    if chinese_font.exists():
        page.insert_font(fontname="SimHei", fontfile=chinese_font)
        page.insert_text((72, 110), "中文测试", fontname="SimHei")
    else:
        page.insert_text((72, 110), "Chinese fallback")
    doc.save(source_pdf)
    doc.close()

    output = tmp_path / "text.docx"
    result = pdf_to_word(source_pdf, output, ocr_enabled=False)

    assert result.output_path == output
    parsed = Document(output)
    text = "\n".join(paragraph.text for paragraph in parsed.paragraphs)
    assert "Hello PDF" in text
    assert ("中文测试" in text) or ("Chinese fallback" in text)


class FakeOcr:
    def recognize(self, image_path: Path) -> list[str]:
        assert image_path.exists()
        return ["识别出的中文", "OCR line"]


def test_pdf_to_word_uses_ocr_for_scanned_pdf(tmp_path: Path) -> None:
    scan_image = tmp_path / "scan.png"
    make_image(scan_image, (255, 255, 255), size=(240, 120))
    source_pdf = tmp_path / "scan.pdf"
    Image.open(scan_image).save(source_pdf)

    output = tmp_path / "scan.docx"
    result = pdf_to_word(source_pdf, output, ocr_enabled=True, ocr_engine=FakeOcr())

    assert result.page_count == 1
    parsed = Document(output)
    text = "\n".join(paragraph.text for paragraph in parsed.paragraphs)
    assert "识别出的中文" in text
    assert "OCR line" in text


class RecordingOcr:
    def __init__(self) -> None:
        self.paths: list[Path] = []

    def recognize(self, image_path: Path) -> list[str]:
        self.paths.append(image_path)
        assert image_path.exists()
        return ["OCR fallback page"]


def test_pdf_to_word_uses_ocr_for_blank_pages_in_mixed_pdf(tmp_path: Path) -> None:
    source_pdf = tmp_path / "mixed.pdf"
    doc = fitz.open()
    text_page = doc.new_page()
    text_page.insert_text((72, 72), "Editable page")
    doc.new_page()
    doc.save(source_pdf)
    doc.close()

    output = tmp_path / "mixed.docx"
    ocr = RecordingOcr()
    result = pdf_to_word(source_pdf, output, ocr_enabled=True, ocr_engine=ocr)

    assert result.page_count == 2
    assert len(ocr.paths) == 1
    parsed = Document(output)
    text = "\n".join(paragraph.text for paragraph in parsed.paragraphs)
    assert "Editable page" in text
    assert "OCR fallback page" in text


def test_ocr_model_dirs_copy_non_ascii_paths_for_paddle(monkeypatch, tmp_path: Path) -> None:
    source = tmp_path / "中文模型" / "ch_PP-OCRv4_det_infer"
    source.mkdir(parents=True)
    (source / "inference.pdmodel").write_bytes(b"model")
    cache_root = tmp_path / "ascii-cache"

    monkeypatch.setattr(converters, "app_data_dir", lambda: cache_root)

    result = converters._ocr_model_dirs_for_paddle({"det_model_dir": source})

    assert result == {"det_model_dir": str(cache_root / "ocr-models" / source.name)}
    assert (cache_root / "ocr-models" / source.name / "inference.pdmodel").exists()


def test_word_to_pdf_reports_missing_libreoffice(tmp_path: Path) -> None:
    source = tmp_path / "sample.docx"
    Document().save(source)

    with pytest.raises(MissingDependencyError, match="LibreOffice"):
        word_to_pdf(source, tmp_path, soffice_path=tmp_path / "missing-soffice.exe")


def test_pdf_to_images_reports_invalid_page_range_with_stable_code(tmp_path: Path) -> None:
    source_pdf = tmp_path / "source.pdf"
    doc = fitz.open()
    doc.new_page()
    doc.save(source_pdf)
    doc.close()

    with pytest.raises(ConversionError) as exc_info:
        pdf_to_images(source_pdf, tmp_path / "pages", pages=[2])

    assert exc_info.value.code == "invalid_pages"
    assert "页码" in exc_info.value.action


def test_pdf_to_word_reports_scanned_pdf_without_ocr_with_stable_code(tmp_path: Path) -> None:
    scan_image = tmp_path / "scan.png"
    make_image(scan_image, (255, 255, 255), size=(240, 120))
    source_pdf = tmp_path / "scan.pdf"
    Image.open(scan_image).save(source_pdf)

    with pytest.raises(ConversionError) as exc_info:
        pdf_to_word(source_pdf, tmp_path / "scan.docx", ocr_enabled=False)

    assert exc_info.value.code == "scanned_pdf_requires_ocr"
    assert "OCR" in exc_info.value.action
