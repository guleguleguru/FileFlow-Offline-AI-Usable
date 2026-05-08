from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
import shutil
import subprocess
import tempfile
from typing import Protocol, Sequence

import fitz
from docx import Document
from docx.enum.section import WD_SECTION
from docx.shared import Pt
from PIL import Image, ImageOps

from offline_converter.dependencies import add_ocr_plugin_paths, find_bundled_ocr_models, find_soffice
from offline_converter.errors import FileFlowError
from offline_converter.logging_utils import app_data_dir


class ConversionError(FileFlowError):
    """Raised when a conversion fails."""


class MissingDependencyError(ConversionError):
    """Raised when an optional conversion engine is unavailable."""


class OcrEngine(Protocol):
    def recognize(self, image_path: Path) -> list[str]:
        ...


@dataclass(frozen=True)
class ConversionResult:
    output_path: Path
    page_count: int = 0
    outputs: tuple[Path, ...] = ()


def image_to_pdf(
    image_paths: Sequence[Path | str],
    output_path: Path | str,
    *,
    quality: int = 90,
    auto_rotate: bool = True,
) -> ConversionResult:
    paths = [_require_file(path) for path in image_paths]
    if not paths:
        raise ConversionError(
            "At least one image is required.",
            code="invalid_input",
            action="请至少选择一张图片后再转换。",
        )

    output = Path(output_path)
    output.parent.mkdir(parents=True, exist_ok=True)

    images: list[Image.Image] = []
    try:
        for path in paths:
            image = Image.open(path)
            if auto_rotate:
                image = ImageOps.exif_transpose(image)
            images.append(_to_pdf_image(image))

        first, rest = images[0], images[1:]
        first.save(
            output,
            "PDF",
            save_all=True,
            append_images=rest,
            quality=max(1, min(100, quality)),
            resolution=100.0,
        )
    finally:
        for image in images:
            image.close()

    return ConversionResult(output_path=output, page_count=len(paths), outputs=(output,))


def pdf_to_images(
    pdf_path: Path | str,
    output_dir: Path | str,
    *,
    image_format: str = "png",
    dpi: int = 150,
    pages: Sequence[int] | None = None,
) -> ConversionResult:
    pdf = _require_file(pdf_path)
    destination = Path(output_dir)
    destination.mkdir(parents=True, exist_ok=True)
    suffix = image_format.lower().lstrip(".")
    if suffix not in {"png", "jpg", "jpeg"}:
        raise ConversionError(
            "Image format must be png, jpg, or jpeg.",
            code="invalid_image_format",
            action="请把图片导出格式设置为 png 或 jpg。",
        )

    exported: list[Path] = []
    with fitz.open(pdf) as document:
        page_numbers = _normalize_pages(pages, document.page_count)
        zoom = dpi / 72
        matrix = fitz.Matrix(zoom, zoom)
        for page_number in page_numbers:
            page = document.load_page(page_number - 1)
            pixmap = page.get_pixmap(matrix=matrix, alpha=False)
            extension = "jpg" if suffix == "jpeg" else suffix
            output = destination / f"{pdf.stem}-page-{page_number:03d}.{extension}"
            pixmap.save(output)
            exported.append(output)

    return ConversionResult(output_path=destination, page_count=len(exported), outputs=tuple(exported))


def pdf_to_word(
    pdf_path: Path | str,
    output_path: Path | str,
    *,
    ocr_enabled: bool = True,
    ocr_engine: OcrEngine | None = None,
    mode: str = "editable",
) -> ConversionResult:
    pdf = _require_file(pdf_path)
    output = Path(output_path)
    output.parent.mkdir(parents=True, exist_ok=True)
    if mode == "visual":
        return _pdf_to_visual_word(pdf, output)
    if mode != "editable":
        raise ConversionError(
            f"Unsupported PDF to Word mode: {mode}",
            code="invalid_pdf_word_mode",
            action="请使用 editable 或 visual。",
        )

    word = Document()
    pages_written = 0
    with fitz.open(pdf) as document:
        text_pages = [_extract_page_text(page) for page in document]
        if not any(page_text.strip() for page_text in text_pages) and not ocr_enabled:
            raise ConversionError(
                "PDF does not contain editable text. Enable OCR for scanned documents.",
                code="scanned_pdf_requires_ocr",
                action="请安装 OCR Addon 并启用 OCR，或改用 visual 模式保留页面外观。",
            )

        engine: OcrEngine | None = None
        temp_dir: tempfile.TemporaryDirectory[str] | None = None
        try:
            for index, page in enumerate(document, start=1):
                page_text = text_pages[index - 1]
                if index > 1:
                    word.add_page_break()
                if page_text.strip():
                    _add_text_to_document(word, page_text)
                elif ocr_enabled:
                    if engine is None:
                        engine = ocr_engine or PaddleOcrEngine()
                    if temp_dir is None:
                        temp_dir = tempfile.TemporaryDirectory(prefix="offline-converter-ocr-")
                    image_path = Path(temp_dir.name) / f"page-{index:03d}.png"
                    pixmap = page.get_pixmap(matrix=fitz.Matrix(2, 2), alpha=False)
                    pixmap.save(image_path)
                    lines = engine.recognize(image_path)
                    for line in lines:
                        if line.strip():
                            word.add_paragraph(line.strip())
                pages_written += 1
        finally:
            if temp_dir is not None:
                temp_dir.cleanup()

    word.save(output)
    return ConversionResult(output_path=output, page_count=pages_written, outputs=(output,))


def word_to_pdf(
    document_path: Path | str,
    output_dir: Path | str,
    *,
    soffice_path: Path | str | None = None,
) -> ConversionResult:
    source = _require_file(document_path)
    destination = Path(output_dir)
    destination.mkdir(parents=True, exist_ok=True)
    soffice = _resolve_soffice(soffice_path)

    completed = subprocess.run(
        [
            str(soffice),
            "--headless",
            "--convert-to",
            "pdf",
            "--outdir",
            str(destination),
            str(source),
        ],
        capture_output=True,
        text=True,
        timeout=180,
        check=False,
    )
    output = destination / f"{source.stem}.pdf"
    if completed.returncode != 0 or not output.exists():
        details = (completed.stderr or completed.stdout or "unknown LibreOffice error").strip()
        raise ConversionError(
            f"LibreOffice failed to convert '{source.name}'.",
            code="word_to_pdf_failed",
            action="请确认 LibreOffice Addon 完整安装，或导出日志查看 LibreOffice 失败原因。",
            detail=details,
        )

    return ConversionResult(output_path=output, page_count=1, outputs=(output,))


class PaddleOcrEngine:
    def __init__(self) -> None:
        add_ocr_plugin_paths()
        try:
            from paddleocr import PaddleOCR
        except ImportError as exc:
            raise MissingDependencyError(
                "PaddleOCR is required for scanned PDF OCR. Install the OCR bundle or run "
                "`python -m pip install paddleocr paddlepaddle`.",
                code="missing_ocr",
                action="Install OCR Addon, or confirm plugins/ocr/python contains a complete PaddleOCR runtime.",
                detail=str(exc),
            ) from exc

        model_dirs = _ocr_model_dirs_for_paddle(find_bundled_ocr_models())
        self._ocr = PaddleOCR(use_angle_cls=True, lang="ch", show_log=False, **model_dirs)

    def recognize(self, image_path: Path) -> list[str]:
        result = self._ocr.ocr(str(image_path), cls=True)
        lines: list[str] = []
        for page_result in result or []:
            for item in page_result or []:
                if len(item) >= 2 and isinstance(item[1], (list, tuple)) and item[1]:
                    text = str(item[1][0]).strip()
                    if text:
                        lines.append(text)
        return lines


def _require_file(path: Path | str) -> Path:
    resolved = Path(path)
    if not resolved.exists() or not resolved.is_file():
        raise ConversionError(
            f"Input file does not exist: {resolved}",
            code="invalid_input",
            action="请选择存在的输入文件，并确认文件类型与转换方式匹配。",
        )
    return resolved


def _to_pdf_image(image: Image.Image) -> Image.Image:
    if image.mode == "RGB":
        return image.copy()
    if image.mode in {"RGBA", "LA"}:
        background = Image.new("RGB", image.size, "white")
        alpha = image.getchannel("A") if "A" in image.getbands() else None
        background.paste(image, mask=alpha)
        return background
    return image.convert("RGB")


def _normalize_pages(pages: Sequence[int] | None, page_count: int) -> list[int]:
    if pages is None:
        return list(range(1, page_count + 1))
    normalized = list(dict.fromkeys(int(page) for page in pages))
    invalid = [page for page in normalized if page < 1 or page > page_count]
    if invalid:
        raise ConversionError(
            f"Page numbers out of range: {invalid}",
            code="invalid_pages",
            action=f"请使用 1 到 {page_count} 之间的页码范围。",
        )
    return normalized


def _extract_page_text(page: fitz.Page) -> str:
    blocks = page.get_text("blocks")
    text_blocks = []
    for block in blocks:
        if len(block) >= 5 and str(block[4]).strip():
            text_blocks.append((float(block[1]), float(block[0]), str(block[4]).strip()))
    text_blocks.sort(key=lambda item: (item[0], item[1]))
    return "\n".join(block[2] for block in text_blocks)


def _add_text_to_document(word: Document, page_text: str) -> None:
    for line in page_text.splitlines():
        cleaned = line.strip()
        if cleaned:
            word.add_paragraph(cleaned)


def _ocr_model_dirs_for_paddle(model_dirs: dict[str, Path]) -> dict[str, str]:
    if not model_dirs:
        return {}
    if all(_is_ascii_path(path) for path in model_dirs.values()):
        return {name: str(path) for name, path in model_dirs.items()}
    cache_root = app_data_dir() / "ocr-models"
    cache_root.mkdir(parents=True, exist_ok=True)
    normalized: dict[str, str] = {}
    for name, source in model_dirs.items():
        target = cache_root / source.name
        if not (target / "inference.pdmodel").exists():
            if target.exists():
                shutil.rmtree(target)
            shutil.copytree(source, target)
        normalized[name] = str(target)
    return normalized


def _is_ascii_path(path: Path) -> bool:
    try:
        str(path).encode("ascii")
    except UnicodeEncodeError:
        return False
    return True


def _pdf_to_visual_word(pdf: Path, output: Path) -> ConversionResult:
    with tempfile.TemporaryDirectory(prefix="fileflow-visual-word-") as temp_dir:
        page_count = _build_visual_docx(pdf, output, Path(temp_dir))
    return ConversionResult(output_path=output, page_count=page_count, outputs=(output,))


def _build_visual_docx(pdf: Path, output: Path, temp_path: Path) -> int:
    word = Document()
    output.parent.mkdir(parents=True, exist_ok=True)
    with fitz.open(pdf) as document:
        for index, page in enumerate(document, start=1):
            rect = page.rect
            section = word.sections[0] if index == 1 else word.add_section(WD_SECTION.NEW_PAGE)
            _configure_visual_section(section, rect.width, rect.height)

            image_path = temp_path / f"page-{index:03d}.png"
            pixmap = page.get_pixmap(matrix=fitz.Matrix(2, 2), alpha=False)
            pixmap.save(image_path)

            paragraph = word.add_paragraph()
            _make_paragraph_tight(paragraph)
            picture_run = paragraph.add_run()
            # A nearly full-page inline image avoids creating trailing blank pages
            # from Word's paragraph marker while preserving the PDF appearance.
            picture_run.add_picture(
                str(image_path),
                width=Pt(float(rect.width)),
                height=Pt(max(1.0, float(rect.height) - 2.0)),
            )
        page_count = document.page_count
    word.save(output)
    return page_count


def _configure_visual_section(section: object, width_pt: float, height_pt: float) -> None:
    section.page_width = Pt(float(width_pt))
    section.page_height = Pt(float(height_pt))
    section.top_margin = Pt(0)
    section.bottom_margin = Pt(0)
    section.left_margin = Pt(0)
    section.right_margin = Pt(0)
    section.header_distance = Pt(0)
    section.footer_distance = Pt(0)


def _make_paragraph_tight(paragraph: object) -> None:
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)


def _resolve_soffice(soffice_path: Path | str | None) -> Path:
    if soffice_path is not None:
        candidate = Path(soffice_path)
        if candidate.exists() and candidate.is_file():
            return candidate
        raise MissingDependencyError(
            f"LibreOffice executable was not found: {candidate}",
            code="missing_libreoffice",
            action="请安装 LibreOffice Addon，或确认 soffice.exe 路径正确。",
        )

    bundled = Path(__file__).resolve().parents[2] / "vendor" / "LibreOffice" / "program" / "soffice.exe"
    candidates = [
        bundled,
        Path("C:/Program Files/LibreOffice/program/soffice.exe"),
        Path("C:/Program Files (x86)/LibreOffice/program/soffice.exe"),
    ]
    discovered = find_soffice()
    if discovered:
        candidates.insert(0, discovered)

    for candidate in candidates:
        if candidate.exists() and candidate.is_file():
            return candidate

    raise MissingDependencyError(
        "LibreOffice is required for Word to PDF conversion. Bundle LibreOffice under "
        "`vendor/LibreOffice` or install LibreOffice locally.",
        code="missing_libreoffice",
        action="请安装 LibreOffice Addon，或安装本机 LibreOffice。",
    )
